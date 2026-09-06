from __future__ import annotations

import hashlib
import os
import sys
import time
import zipfile

import pytest

from ollama_code import document_library as library
from ollama_code.document_extract import ExtractionError, extract_office
from ollama_code.knowledge import KnowledgeStore, format_search_results


class ManualCoordinator:
    def register(self, _store):
        pass


@pytest.fixture
def store(tmp_path, monkeypatch):
    monkeypatch.setattr(library, "_coordinator", lambda: ManualCoordinator())
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    KnowledgeStore(str(workspace)).configure(documents_enabled=True)
    return library.DocumentStore(str(workspace), start_worker=False)


def finish(store, job):
    store._execute(job["id"])
    final = store.job(job["id"])
    assert final["state"] in {"ready", "partial"}, final
    return store.result(job["id"])


def test_document_opt_in_migration_preserves_existing_text_index(tmp_path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "a.py").write_text("old_python_marker = True")
    (workspace / "table.csv").write_text("account,amount\nsearchable,41\n")
    knowledge = KnowledgeStore(str(workspace))
    assert knowledge.settings()["documents_enabled"] is False
    knowledge.reindex()
    assert knowledge.search("old_python_marker")
    assert knowledge.search("searchable") == []
    db = knowledge._connect()
    with db:
        db.execute("ALTER TABLE settings DROP COLUMN documents_enabled")
        db.execute("ALTER TABLE documents DROP COLUMN format")
        db.execute("ALTER TABLE chunks DROP COLUMN locator_json")
    migrated = KnowledgeStore(str(workspace))
    assert migrated.settings()["documents_enabled"] is False
    assert migrated.search("old_python_marker")[0]["line_start"] == 1


def test_csv_persistent_job_indexes_hash_and_cell_citations(store):
    path = store.root / "quarterly.csv"
    path.write_text('Account,Amount\n"North, Inc",125\n', encoding="utf-8")
    job = store.submit(path.name)
    result = finish(store, job)
    assert result["content_hash"] == hashlib.sha256(path.read_bytes()).hexdigest()
    assert result["segments"][1]["locator"] == {"kind": "sheet", "sheet": "quarterly", "cell_range": "A2:B2"}
    knowledge = KnowledgeStore(str(store.root))
    found = knowledge.search("North")
    assert found[0]["content_hash"] == result["content_hash"]
    assert found[0]["line_start"] == 0
    assert found[0]["locator"]["cell_range"] == "A2:B2"
    citation = format_search_results(found)
    assert "locus-workspace://open/quarterly.csv?" in citation
    assert "hash=" in citation and "quarterly!A2:B2" in citation
    assert store.submit(path.name)["id"] == job["id"]
    assert store.documents()["total"] == 1


def test_unchanged_reconciliation_does_not_snapshot_or_retry_failed_sources(store, monkeypatch):
    ready = store.root / "ready.csv"
    ready.write_text("indexedfact\n")
    ready_job = store.submit(ready.name, ocr_mode="always", ocr_languages=["en-US"])
    finish(store, ready_job)
    broken = store.root / "broken.pdf"
    broken.write_bytes(b"broken PDF")
    monkeypatch.delenv("LOCUS_DOCUMENT_EXTRACTOR_PATH", raising=False)
    failed = store.submit(broken.name)
    store._execute(failed["id"])
    assert store.job(failed["id"])["state"] == "failed"

    original = store._submit_source
    snapshots = []

    def snapshot(*args, **kwargs):
        snapshots.append(args[1])
        return original(*args, **kwargs)

    monkeypatch.setattr(store, "_submit_source", snapshot)
    for _ in range(3):
        store.reconcile()
    assert snapshots == []
    # Discovery also uses the fast path, including across store instances.
    KnowledgeStore(str(store.root)).reindex()
    assert len(store.jobs()) == 2
    assert store.job(ready_job["id"])["options"]["ocr_mode"] == "always"
    retried = store.submit(broken.name)
    assert retried["id"] != failed["id"]
    assert snapshots == [broken.name]


def test_ctime_fingerprint_detects_same_size_edits_with_restored_mtime(store):
    source = store.root / "changed.csv"
    source.write_text("firstvalue\n")
    first = store.submit(source.name)
    finish(store, first)
    previous = source.stat()
    source.write_text("othervalue\n")
    os.utime(source, ns=(previous.st_atime_ns, previous.st_mtime_ns))
    assert source.stat().st_size == previous.st_size
    store.reconcile()
    current = store.documents()["documents"][0]["job_id"]
    assert current != first["id"]
    finish(store, store.job(current))
    assert KnowledgeStore(str(store.root)).search("othervalue")
    assert KnowledgeStore(str(store.root)).search("firstvalue") == []


def test_extraction_cache_reuses_crosspath_hash_and_repairs_cleared_index(store, monkeypatch):
    first_path = store.root / "original.csv"
    first_path.write_text("sharedfact,42\n")
    first = store.submit(first_path.name)
    finish(store, first)
    calls = []
    original = library._extract

    def extract(store, job):
        calls.append(job["id"])
        return original(store, job)

    monkeypatch.setattr(library, "_extract", extract)
    duplicate = store.root / "duplicate.csv"
    duplicate.write_bytes(first_path.read_bytes())
    second = store.submit(duplicate.name)
    result = finish(store, second)
    assert calls == []
    assert result["segments"][0]["locator"]["sheet"] == "duplicate"
    assert store.result(first["id"])["segments"][0]["locator"]["sheet"] == "original"
    knowledge = KnowledgeStore(str(store.root))
    assert {row["path"] for row in knowledge.search("sharedfact")} == {first_path.name, duplicate.name}
    knowledge.remove_document_chunks(duplicate.name)
    store.reconcile()
    repaired = next(row for row in store.documents()["documents"] if row["path"] == duplicate.name)
    assert repaired["job_id"] != second["id"]
    finish(store, store.job(repaired["job_id"]))
    assert calls == []
    assert any(row["path"] == duplicate.name for row in knowledge.search("sharedfact"))
    # Different options must never borrow an incompatible extraction.
    forced = store.submit(duplicate.name, ocr_mode="always")
    finish(store, forced)
    assert calls == [forced["id"]]
    monkeypatch.setattr(library, "EXTRACTOR_VERSION", "documents-next-version")
    newer = store.submit(duplicate.name, automatic=True)
    finish(store, newer)
    assert calls == [forced["id"], newer["id"]]


def test_expired_temporary_cache_and_disabled_publication_are_not_reused(store, monkeypatch):
    source = store.root / "temporary.csv"
    source.write_text("temporaryfact\n")
    temporary = store.submit(source.name, persistent=False)
    finish(store, temporary)
    with store._connect() as db:
        db.execute("UPDATE document_jobs SET expires_at=0 WHERE id=?", (temporary["id"],))
    calls = []
    original = library._extract

    def extract(store, job):
        calls.append(job["id"])
        return original(store, job)

    monkeypatch.setattr(library, "_extract", extract)
    persistent = store.submit(source.name)
    finish(store, persistent)
    assert calls == [persistent["id"]]
    duplicate = store.root / "duplicate.csv"
    duplicate.write_bytes(source.read_bytes())
    queued = store.submit(duplicate.name)
    knowledge = KnowledgeStore(str(store.root))
    knowledge.configure(documents_enabled=False)
    store._execute(queued["id"])
    assert store.job(queued["id"])["state"] == "failed"
    assert knowledge.search("temporaryfact") == []
    assert calls == [persistent["id"]]


def test_disabling_queued_extraction_allows_reenable_to_rediscover(store):
    source = store.root / "resume.csv"
    source.write_text("resumevalue\n")
    original = store.submit(source.name)
    knowledge = KnowledgeStore(str(store.root))
    knowledge.configure(documents_enabled=False)
    store.cancel_persistent()
    knowledge.configure(documents_enabled=True)
    store.reconcile()
    replacement = store.documents()["documents"][0]["job_id"]
    assert replacement != original["id"]
    finish(store, store.job(replacement))
    assert knowledge.search("resumevalue")


def test_ephemeral_results_never_enter_index_and_expire(store):
    source = store.root / "ephemeral.tsv"
    source.write_text("privateone\tprivatetwo\n")
    job = store.submit(source.name, persistent=False)
    result = finish(store, job)
    assert result["segments"]
    assert job["expires_at"] - job["created_at"] == library.TEMPORARY_LIFETIME
    assert store.documents()["total"] == 0
    assert KnowledgeStore(str(store.root)).search("privateone") == []
    with store._connect() as db:
        db.execute("UPDATE document_jobs SET expires_at=0 WHERE id=?", (job["id"],))
    store.cleanup_expired()
    with pytest.raises(library.DocumentError, match="expired"):
        store.job(job["id"])
    assert not (store.job_directory / job["id"]).exists()


def test_cancel_and_source_generation_cannot_publish_late_results(store, monkeypatch):
    source = store.root / "revision.csv"
    source.write_text("firstvalue\n")
    first = store.submit(source.name)
    source.write_text("secondvalue\n")
    second = store.submit(source.name)
    assert store.job(first["id"])["state"] == "cancelled"
    store._execute(first["id"])
    assert KnowledgeStore(str(store.root)).search("firstvalue") == []
    finish(store, second)
    assert KnowledgeStore(str(store.root)).search("secondvalue")
    source.write_text("thirdvalue\n")
    third = store.submit(source.name)
    original = library._extract

    def late_result(store, job):
        result = original(store, job)
        store.cancel(job["id"])
        return result

    monkeypatch.setattr(library, "_extract", late_result)
    store._execute(third["id"])
    assert store.job(third["id"])["state"] == "cancelled"
    assert KnowledgeStore(str(store.root)).search("thirdvalue") == []


def test_source_changes_during_parse_fail_without_stale_evidence(store, monkeypatch):
    path = store.root / "changing.csv"
    path.write_text("oldvalue\n")
    job = store.submit(path.name)
    original = library._extract

    def change_source(store, current):
        result = original(store, current)
        path.write_text("newvalue\n")
        return result

    monkeypatch.setattr(library, "_extract", change_source)
    store._execute(job["id"])
    assert store.job(job["id"])["state"] == "failed"
    assert "changed" in store.job(job["id"])["error"]
    assert KnowledgeStore(str(store.root)).search("oldvalue") == []


def test_excluding_deleting_disabling_and_reenabling_remove_stale_search(store):
    path = store.root / "record.csv"
    path.write_text("valuablefact\n")
    job = store.submit(path.name)
    finish(store, job)
    store.exclude(job["document_id"], True)
    assert KnowledgeStore(str(store.root)).search("valuablefact") == []
    assert path.exists()
    included = store.exclude(job["document_id"], False)
    finish(store, store.job(included["job_id"]))
    knowledge = KnowledgeStore(str(store.root))
    knowledge.configure(documents_enabled=False)
    assert knowledge.search("valuablefact") == []
    knowledge.configure(documents_enabled=True)
    fresh = store.submit(path.name)
    assert fresh["id"] != included["job_id"]
    finish(store, fresh)
    path.unlink()
    store.reconcile()
    assert knowledge.search("valuablefact") == []
    assert store.documents()["documents"][0]["status"] == "unavailable"


def test_paths_symlinks_and_quota_rejected_before_snapshot(store, tmp_path):
    outside = tmp_path / "outside.csv"
    outside.write_text("forbidden")
    (store.root / "link.csv").symlink_to(outside)
    for value in ("../outside.csv", str(outside), "link.csv"):
        with pytest.raises(library.DocumentError):
            store.submit(value)
    huge = store.root / "huge.csv"
    with huge.open("wb") as handle:
        handle.truncate(library.MAX_SOURCE_BYTES + 1)
    with pytest.raises(library.DocumentError, match="100 MB"):
        store.submit(huge.name)


def test_upload_collision_preserves_existing_and_uses_visible_folder(store, tmp_path):
    source = tmp_path / "incoming.csv"
    source.write_text("newfact\n")
    folder = store.root / "Locus Documents"
    folder.mkdir()
    (folder / "incoming.csv").write_text("original\n")
    job = store.submit_upload(source, "incoming.csv", persistent=True)
    assert job["path"] == "Locus Documents/incoming 2.csv"
    assert (folder / "incoming.csv").read_text() == "original\n"


def test_docx_nested_table_heading_and_paragraph_locations(tmp_path):
    from docx import Document

    path = tmp_path / "brief.docx"
    doc = Document()
    doc.add_heading("Launch decision", 1)
    doc.add_paragraph("Use the northern route.")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Owner is Morgan"
    nested = table.cell(0, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "Nested evidence"
    doc.save(path)
    result = extract_office(path, "docx")
    text = "\n".join(item["text"] for item in result["segments"])
    assert "Nested evidence" in text
    assert result["segments"][1]["locator"]["heading"] == "Launch decision"
    assert result["segments"][1]["locator"]["paragraph_start"] == 2
    assert all("page" not in item["locator"] for item in result["segments"])


def test_xlsx_uses_saved_values_addresses_and_ignores_hidden_sheets(tmp_path):
    from openpyxl import Workbook

    path = tmp_path / "numbers.xlsx"
    book = Workbook()
    sheet = book.active
    sheet.title = "North Region"
    sheet.append(["Sales", 150, "=B1*2"])
    hidden = book.create_sheet("Secret")
    hidden.sheet_state = "hidden"
    hidden["A1"] = "hiddenmarker"
    book.save(path)
    result = extract_office(path, "xlsx")
    assert result["segments"][0]["locator"]["sheet"] == "North Region"
    assert result["segments"][0]["locator"]["cell_range"] == "A1:B1"
    text = " ".join(item["text"] for item in result["segments"])
    assert "150" in text and "=B1*2" not in text and "hiddenmarker" not in text
    assert any("saved values" in warning for warning in result["warnings"])


def test_damaged_office_entity_archive_and_text_cell_limits(tmp_path, monkeypatch):
    from ollama_code import document_extract as parser

    path = tmp_path / "invalid.docx"
    path.write_text("not a zip")
    with pytest.raises(ExtractionError):
        extract_office(path, "docx")
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("word/document.xml", '<!DOCTYPE foo [<!ENTITY x "secret">]><root/>')
    with pytest.raises(ExtractionError, match="entity"):
        extract_office(path, "docx")
    table = tmp_path / "table.csv"
    table.write_text("a,b,c\n1,2,3\n")
    monkeypatch.setattr(parser, "MAX_CELLS", 3)
    result = extract_office(table, "csv")
    assert result["truncated"] and len(result["segments"]) == 1
    monkeypatch.setattr(parser, "MAX_TEXT_BYTES", 5)
    result = extract_office(table, "csv")
    assert result["truncated"]
    assert sum(len(item["text"].encode()) for item in result["segments"]) <= 5


def test_crashed_owner_requeues_but_live_owner_is_not_stolen(store):
    path = store.root / "queued.csv"
    path.write_text("fact\n")
    job = store.submit(path.name)
    with store._connect() as db:
        db.execute("UPDATE document_jobs SET state='running',owner_pid=? WHERE id=?", (os.getpid(), job["id"]))
    assert store._pending() == []
    with store._connect() as db:
        db.execute("UPDATE document_jobs SET owner_pid=2147483647 WHERE id=?", (job["id"],))
    assert store._pending() == [job["id"]]


def test_shutdown_retains_snapshot_and_queues_interrupted_job(store, monkeypatch):
    source = store.root / "resume.csv"
    source.write_text("durablemarker\n")
    job = store.submit(source.name)

    def interrupted(*_):
        raise library._JobInterrupted

    monkeypatch.setattr(library, "_extract", interrupted)
    store._execute(job["id"])
    assert store.job(job["id"])["state"] == "queued"
    assert (store.job_directory / job["id"] / "source").exists()
    assert store.documents()["documents"][0]["status"] == "queued"


def test_process_locks_enforce_one_workspace_and_two_global(store, tmp_path):
    # Instantiate without starting the coordinator's scheduler.
    coordinator = object.__new__(library._Coordinator)
    first = coordinator._locks(store)
    assert first
    try:
        assert coordinator._locks(store) is None
        other = tmp_path / "other"
        other.mkdir()
        second_store = library.DocumentStore(str(other), start_worker=False)
        second = coordinator._locks(second_store)
        assert second
        try:
            third = tmp_path / "third"
            third.mkdir()
            third_store = library.DocumentStore(str(third), start_worker=False)
            assert coordinator._locks(third_store) is None
        finally:
            for handle in second:
                handle.close()
    finally:
        for handle in first:
            handle.close()


def test_pdf_helper_unavailable_is_explicit_failure(store, monkeypatch):
    monkeypatch.delenv("LOCUS_DOCUMENT_EXTRACTOR_PATH", raising=False)
    path = store.root / "unreadable.pdf"
    path.write_bytes(b"%PDF-1.4\n")
    job = store.submit(path.name)
    store._execute(job["id"])
    assert store.job(job["id"])["state"] == "failed"
    assert "bundled Locus document helper" in store.job(job["id"])["error"]


def test_timeout_kills_helper_and_does_not_publish_fake_success(store, monkeypatch, tmp_path):
    helper = tmp_path / "slow-helper"
    helper.write_text(f"#!{sys.executable}\nimport time\ntime.sleep(30)\n")
    helper.chmod(0o700)
    monkeypatch.setenv("LOCUS_DOCUMENT_EXTRACTOR_PATH", str(helper))
    monkeypatch.setattr(library, "DOCUMENT_DEADLINE", .1)
    path = store.root / "slow.pdf"
    path.write_bytes(b"%PDF-1.4\n")
    job = store.submit(path.name)
    started = time.monotonic()
    store._execute(job["id"])
    assert time.monotonic() - started < 4
    assert store.job(job["id"])["state"] == "failed"
    assert not store.job(job["id"])["result_available"]


def test_authenticated_http_document_contract_and_upload_scope(store, monkeypatch):
    from types import SimpleNamespace

    from fastapi.testclient import TestClient

    from ollama_code import server

    service = SimpleNamespace(core=SimpleNamespace(workspace_root=str(store.root), cwd=str(store.root)))
    client = TestClient(server.create_app(chat_service=service, auth_token="local-token"))
    assert client.get("/api/documents").status_code == 401
    headers = {"x-locus-token": "local-token"}
    assert client.get("/api/documents", headers={**headers, "origin": "https://example.test"}).status_code == 403
    response = client.post("/api/document-jobs/upload", params={"workspace": str(store.root), "filename": "http.csv", "persistent": False}, content=b"httpmarker,55\n", headers=headers)
    assert response.status_code == 200, response.text
    job = response.json()["job"]
    assert not job["persistent"]
    queried = client.get(f"/api/document-jobs/{job['id']}", headers=headers)
    assert queried.json()["job"]["state"] == "queued"
    store._execute(job["id"])
    result = client.get(f"/api/document-jobs/{job['id']}/result", headers=headers)
    assert result.status_code == 200, result.text
    assert "httpmarker" in result.json()["segments"][0]["text"]
    assert client.get("/api/documents", headers=headers).json()["total"] == 0
    assert client.post("/api/document-jobs/upload", params={"filename": "too-large.pdf"}, content=b"x", headers={**headers, "content-length": str(library.MAX_SOURCE_BYTES + 1)}).status_code == 413
    assert client.post("/api/knowledge/settings", json={}, headers={**headers, "content-length": str(server.MAX_HTTP_BODY_BYTES + 1)}).status_code == 413
    # Chunked bodies are checked as they arrive, not just by Content-Length.
    monkeypatch.setattr("ollama_code.api.knowledge.MAX_SOURCE_BYTES", 5)
    rejected = client.post("/api/document-jobs/upload", params={"filename": "chunked.csv"}, content=iter([b"abc", b"def"]), headers=headers)
    assert rejected.status_code == 413
    client.close()


def test_catalog_filename_search_filters_before_pagination_and_counts_all_matches(store):
    from types import SimpleNamespace

    from fastapi.testclient import TestClient

    from ollama_code import server

    # Older documents sit beyond the initial page, including failed documents
    # whose text cannot appear in content search. Names must remain searchable.
    rows = [(str(index), f"new-{index}.csv", "ready", float(index + 10)) for index in range(105)]
    rows += [("older-a", "Archive/Übersicht 100%.csv", "failed", 1.0),
             ("older-b", "Archive/Übersicht 100% details.csv", "excluded", 0.0)]
    with store._connect() as db:
        db.executemany("""INSERT INTO library_documents(id,path,format,content_hash,size,status,updated_at)
            VALUES(?,?,'csv','hash',1,?,?)""", rows)
    service = SimpleNamespace(core=SimpleNamespace(workspace_root=str(store.root), cwd=str(store.root)))
    client = TestClient(server.create_app(chat_service=service, auth_token="local-token"))
    headers = {"x-locus-token": "local-token"}
    initial = client.get("/api/documents", headers=headers).json()
    assert initial["total"] == 107
    assert len(initial["documents"]) == 100
    assert not any(row["id"].startswith("older") for row in initial["documents"])
    query = {"query": " ÜBERSICHT 100% ", "limit": 1}
    first = client.get("/api/documents", params=query, headers=headers).json()
    assert first["total"] == 2
    assert [row["id"] for row in first["documents"]] == ["older-a"]
    second = client.get("/api/documents", params={**query, "offset": 1}, headers=headers).json()
    assert second["total"] == 2
    assert [row["id"] for row in second["documents"]] == ["older-b"]
    assert client.get("/api/documents", params={"query": "Archive/"}, headers=headers).json()["total"] == 2
    assert client.get("/api/documents", params={"query": "_"}, headers=headers).json()["total"] == 0
    assert client.get("/api/documents", params={"query": "' OR 1=1"}, headers=headers).json()["total"] == 0
    assert client.get("/api/documents", params={"query": "z" * 513}, headers=headers).status_code == 422
    client.close()
